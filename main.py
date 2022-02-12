from bs4 import BeautifulSoup
import requests as request
from docx import Document

document = Document()
document.add_heading('Thanks for use', 0)

novelChapters = []
curChapter = 1646
countChapters = 300
link = ""
i = 0
link = input("Insert novel link. Reference: https://boxnovel.com/novel/tales-of-herding-gods/\n")

curChapter = int(input("Insert chapter where we start. Reference: 1646\n"))

countChapters = int(input("Insert count chapters to add(can be infinity). Reference: 1000\n"))

try:
    while countChapters > 0:
        rHtml = request.get(f"{link}/chapter-{curChapter}/")
        soup = BeautifulSoup(rHtml.text, 'html.parser')
        textContainer = soup.find_all("div", {"class": "text-left"})
        novelChapters.append(textContainer[0].text)
        document.add_paragraph(novelChapters[i])
        curChapter += 1
        i += 1
        countChapters -= 1
        print(f"{curChapter} added to doc")
        document.save('NotTranslated.docx')
except IndexError:
    print("End of chapters")

print(f"All chapters added")
input()
